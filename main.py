import streamlit as st 
from bs4 import BeautifulSoup
import requests
import re 
from docx import Document
from io import BytesIO


st.set_page_config("Web Scraper y Generador de word")

def scraper():
    st.title("👇 :orange[Web scraper a Wikipedia] 😎")
    st.divider()
    tabs = st.tabs(['**Como usar**', '**Descarga word**', '**Desarrollador**'])

    with tabs[0]:
        st.write(":green[1]. Escribe lo que deseas buscar en el campo de texto")
        st.write(":green[2]. Presiona Enter y te traera lo buscado")
        st.write(":green[3]. Puedes descargar el resultado en formato word")
        st.write(":green[4]. Si la cantidad de parrafos :red[No] salen al puesto es porque es toda la informacíon que hay sobre el tema")
        
    
    with tabs[1]:
        st.write(":green[1]. Generar archivo word")
        st.write(":green[2]. Haz clic en el botón de descarga")
        st.write(":green[3]. Se descargará un archivo word con la información")
        st.write(":green[4]. Abrir el archivo word con cualquier programa de word")

    with tabs[2]:
        st.write("🙌 Hans Saldías - Analista Programador 😎")
        st.write("Espero que les sirva 😘")
        st.write("👍🤗🤗😘")

    st.divider()
 


    search = st.text_input("Ingrese lo que desea buscar", placeholder="Ingrese lo que busca - presione (Enter)")
    parrafos = st.number_input("Ingrese la cantidad de parrafos en numero **:red[ENTERO]**", step=1, min_value=1)
    if search:
        base_url = "https://es.wikipedia.org/wiki/"
        url = base_url + search.replace(" ", "_")
        st.write("URL generada " + url)
        try:
            response = requests.get(url)
            response.raise_for_status()
            st.write("**Respuesta del servidor Aceptada**")
            soup = BeautifulSoup(response.text, 'html.parser')
            titulo = soup.find('th').text
            st.write(f"**:green[titulo de articulos] :red[{titulo}]**")
            word_text = []
            paragram = soup.find_all('p')
            for p in paragram[:parrafos]:
                text = re.sub(r'\[d+\]', '', p.text)
                word_text.append(text)
                st.write(text)

            if st.button("Generar archivo word", type='primary', icon=":material/article:"):
                doc = Document()
                doc.add_heading(titulo, level=1)
                for i in word_text:
                    doc.add_paragraph(i)
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                st.download_button(
                    label="Descargar documento",
                    icon=":material/download:",
                    data=buffer,
                    file_name=f"{titulo}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
        except requests.exceptions.RequestException as e:
            st.error(f"❌ Error de conexion o no encuentra lo buscado, intente de otra forma al escribir. 😒")

        except Exception as e:
            st.error(f"Error al analizar el html {e}")


scraper()